import hashlib
import re
import unittest
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
CONTENT_PATH = ROOT / "docs" / "cv-content.md"
INVENTORY_PATH = ROOT / "docs" / "publications-inventory.md"
DOCX_PATH = ROOT / "files" / "Guo-Lu-CV.docx"
PDF_PATH = ROOT / "files" / "Guo-Lu-CV.pdf"
LEGACY_PDF_PATH = ROOT / "paper" / "GuoLu.pdf"
EXPECTED_PROFILE_LINKS = {
    "https://guolusjtu.github.io/guoluhomepage/",
    "https://scholar.google.com/citations?user=R9iwlJcAAAAJ&hl=en",
    "https://github.com/GuoLusjtu?tab=repositories",
    "https://www.linkedin.com/in/guo-lu-118a6592/",
}

EXPECTED_TITLES = [
    "Next-frame decoding for ultra-low-bitrate image compression with video diffusion priors",
    "Every Packet Counts: Dispersing Information for Loss-Resilient Learned Image Compression",
    "Diff-VF: Training-free High-quality Long Video Generation via Diffusion Model",
    "Large language model for lossless image compression with visual prompts",
    "Generative Video Communications: Concepts, Key Technologies, and Future Research Trends",
    "Unified spatiotemporal token compression for video-llms at ultra-low retention",
    "Adaptive Learned Image Compression with Graph Neural Networks",
    "Content-aware mamba for learned image compression",
    "Knowledge Distillation for Learned Image Compression",
    "Controllable Distortion-Perception Tradeoff Through Latent Diffusion for Neural Image Compression",
    "SMC++: Masked Learning of Unsupervised Video Semantic Compression",
    "Image Quality Assessment: From Human to Machine Preference",
    "Towards Trustworthy Multimodal Moderation via Policy-Aligned Reasoning and Hierarchical Labeling",
    "Implicit-Explicit Integrated Representations for Multi-View Video Compression",
    "VARFVV: View-Adaptive Real-Time Interactive Free-View Video Streaming With Edge Computing",
    "A Coding Framework and Benchmark Towards Low-Bitrate Video Understanding",
    "Neural rate control for learned video compression",
    "Preprocessing Enhanced Image Compression for Machine Vision",
    "Rate-aware Compression for NeRF-based Volumetric Video",
    "FVC: A New Framework towards Deep Video Compression in Feature Space",
    "DVC: An End-To-End Deep Video Compression Framework",
]

GRANT_IDENTIFIER_PATTERN = re.compile(
    r"\b(?:grant|project)\s*(?:no\.?|number|id)\s*[:#]?\s*[A-Z0-9][A-Z0-9-]{3,}"
    r"|\bNSFC\s*(?:no\.?\s*)?[:#]?\s*\d{6,}\b",
    flags=re.IGNORECASE,
)
FUNDING_AMOUNT_PATTERN = re.compile(
    r"(?:\b(?:CNY|RMB|USD)\b|US\$|[$¥])\s*\d"
    r"|\b\d+(?:\.\d+)?\s*(?:million|billion|thousand)\s*(?:yuan|CNY|RMB|USD)\b"
    r"|\b(?:\d{1,3}(?:,\d{3})+|\d+(?:\.\d+)?)\s*(?:yuan|CNY|RMB|USD)\b"
    r"|\d+(?:\.\d+)?\s*万(?:元)?",
    flags=re.IGNORECASE,
)
CORRESPONDING_AUTHOR_PATTERN = re.compile(
    r"corresponding[- ]author|Guo Lu\s*(?:\([*†‡#]\)|[*†‡#])|[*†‡]\s*Guo Lu",
    flags=re.IGNORECASE,
)


def parse_markdown_table(text, heading):
    match = re.search(
        rf"^## {re.escape(heading)}\s*$\n(.*?)(?=^## |\Z)",
        text,
        flags=re.MULTILINE | re.DOTALL,
    )
    if not match:
        return []
    rows = []
    for line in match.group(1).splitlines():
        if not line.startswith("|"):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and not all(re.fullmatch(r":?-+:?", cell) for cell in cells):
            rows.append(cells)
    return rows[1:]


def parse_inventory():
    lines = INVENTORY_PATH.read_text(encoding="utf-8").splitlines()
    header_index = next(i for i, line in enumerate(lines) if line.startswith("| Scholar title |"))
    headers = [cell.strip() for cell in lines[header_index].strip("|").split("|")]
    records = []
    for line in lines[header_index + 2 :]:
        if not line.startswith("|"):
            break
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        records.append(dict(zip(headers, cells)))
    return records


class CvContentTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.content = CONTENT_PATH.read_text(encoding="utf-8")
        cls.publication_rows = parse_markdown_table(cls.content, "Selected Publications")
        cls.evidence_rows = parse_markdown_table(cls.content, "Evidence")
        cls.inventory = parse_inventory()

    def test_exact_21_selected_publications_resolve_once_to_included_inventory(self):
        self.assertEqual(21, len(self.publication_rows))
        actual_titles = [row[0] for row in self.publication_rows]
        self.assertEqual(EXPECTED_TITLES, actual_titles)
        included = [row for row in self.inventory if row["Status"] == "include"]
        for title in actual_titles:
            matches = [row for row in included if row["Canonical title"] == title]
            self.assertEqual(1, len(matches), title)

    def test_selected_publication_metadata_exactly_matches_inventory(self):
        self.assertTrue(self.publication_rows)
        self.assertTrue(all(len(row) == 6 for row in self.publication_rows))
        included = {
            row["Canonical title"]: row
            for row in self.inventory
            if row["Status"] == "include"
        }
        for title, authors, venue, year, link, publication_type in self.publication_rows:
            source = included[title]
            self.assertEqual(source["Display authors"], authors, title)
            self.assertEqual(source["Venue"], venue, title)
            self.assertEqual(source["Year"], year, title)
            self.assertEqual(source["Destination"], link, title)
            self.assertEqual(source["Type"], publication_type, title)

    def test_non_publication_facts_have_compact_provenance(self):
        body, separator, evidence_and_definitions = self.content.partition("\n## Evidence\n")
        self.assertTrue(separator)
        body_ids = set(re.findall(r"\[\^(F\d+)\]", body))
        definition_ids = set(
            re.findall(r"^\[\^(F\d+)\]:", evidence_and_definitions, flags=re.MULTILINE)
        )
        self.assertGreaterEqual(len(body_ids), 1)
        evidence_ids = {row[0] for row in self.evidence_rows}
        self.assertEqual(body_ids, evidence_ids)
        self.assertEqual(body_ids, definition_ids)
        for row in self.evidence_rows:
            self.assertEqual(4, len(row))
            self.assertRegex(row[0], r"^F\d+$")
            self.assertTrue(row[1] and row[2])
            self.assertRegex(row[3], r"^(index\.html|https://icisee\.sjtu\.edu\.cn/jiaoshiml/luguo\.html|https://guolusjtu\.github\.io/guoluhomepage/paper/GuoLu\.pdf)(?: \(.+\))?$")

    def test_content_has_no_placeholders_or_unsupported_claim_markers(self):
        forbidden = [
            "TODO", "TBD", "PLACEHOLDER", "grant no.", "grant number",
            "CNY", "RMB", "¥", "$", "corresponding author", "* Guo Lu",
            "†", "‡",
        ]
        lowered = self.content.lower()
        for token in forbidden:
            self.assertNotIn(token.lower(), lowered)
        self.assertIsNone(GRANT_IDENTIFIER_PATTERN.search(self.content))
        self.assertIsNone(FUNDING_AMOUNT_PATTERN.search(self.content))
        self.assertIsNone(CORRESPONDING_AUTHOR_PATTERN.search(self.content))

    def test_negative_claim_patterns_cover_common_unsupported_forms(self):
        for sample in (
            "Grant No. 62300001",
            "Project ID ABCD-1234",
            "NSFC No. 62300001",
            "NSFC 62300001",
        ):
            self.assertRegex(sample, GRANT_IDENTIFIER_PATTERN)
        for sample in (
            "RMB 500000",
            "US$250000",
            "2.5 million yuan",
            "30万元",
            "500000 RMB",
            "500,000 yuan",
            "50万",
        ):
            self.assertRegex(sample, FUNDING_AMOUNT_PATTERN)
        for sample in ("Guo Lu*", "Guo Lu (†)", "‡ Guo Lu", "corresponding-author"):
            self.assertRegex(sample, CORRESPONDING_AUTHOR_PATTERN)
        for legitimate_date in (
            "2022–2024",
            "January 2025–December 2028",
            "NSFC General Program, January 2025–December 2028",
        ):
            self.assertIsNone(GRANT_IDENTIFIER_PATTERN.search(legitimate_date))
            self.assertIsNone(FUNDING_AMOUNT_PATTERN.search(legitimate_date))

    def test_research_interests_preserve_exact_homepage_wording(self):
        section = re.search(
            r"^## Research Interests\s*$\n(.*?)(?=^## )",
            self.content,
            flags=re.MULTILINE | re.DOTALL,
        ).group(1)
        self.assertIn(
            "Video coding, multimedia processing, and efficient multimodal large language models.[^F04]",
            section,
        )
        self.assertNotIn("Multimedia communications", section)

    def test_ambiguous_lac_role_is_not_normalized(self):
        self.assertNotRegex(self.content, r"Technical Program (?:Co-)?Chair")

    def test_generated_cv_files_exist_and_public_pdfs_are_identical(self):
        for path in (DOCX_PATH, PDF_PATH, LEGACY_PDF_PATH):
            self.assertTrue(path.is_file(), path)
            self.assertGreater(path.stat().st_size, 1000, path)
        digest = lambda path: hashlib.sha256(path.read_bytes()).hexdigest()
        self.assertEqual(digest(PDF_PATH), digest(LEGACY_PDF_PATH))

    def test_docx_uses_native_title_headings_hyperlinks_and_page_field(self):
        document = Document(DOCX_PATH)
        self.assertEqual("Title", document.paragraphs[0].style.name)
        styles = [p.style.name for p in document.paragraphs]
        self.assertIn("Heading 1", styles)
        self.assertIn("Heading 2", styles)
        with zipfile.ZipFile(DOCX_PATH) as package:
            document_xml = package.read("word/document.xml").decode("utf-8")
            footer_xml = "".join(
                package.read(name).decode("utf-8")
                for name in package.namelist()
                if name.startswith("word/footer") and name.endswith(".xml")
            )
            rels_xml = package.read("word/_rels/document.xml.rels").decode("utf-8")
        self.assertIn("PAGE", footer_xml)
        self.assertIn("w:hyperlink", document_xml)
        for url in EXPECTED_PROFILE_LINKS:
            self.assertIn(url.replace("&", "&amp;"), rels_xml)

    def test_pdf_is_four_to_six_pages_searchable_and_has_exact_profile_links(self):
        reader = PdfReader(PDF_PATH)
        self.assertGreaterEqual(len(reader.pages), 4)
        self.assertLessEqual(len(reader.pages), 6)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        self.assertIn("Guo Lu", text)
        self.assertIn("Selected Publications", text)
        self.assertGreater(len(text), 5000)
        uris = []
        for page in reader.pages:
            for annotation_ref in page.get("/Annots", []):
                annotation = annotation_ref.get_object()
                action = annotation.get("/A")
                if action and action.get("/URI"):
                    uris.append(str(action["/URI"]))
        for url in EXPECTED_PROFILE_LINKS:
            self.assertEqual(1, uris.count(url), url)

    def test_generated_cv_contains_each_selected_paper_exactly_once_and_no_placeholders(self):
        document = Document(DOCX_PATH)
        docx_text = "\n".join(p.text for p in document.paragraphs)
        pdf_text = "\n".join(
            page.extract_text() or "" for page in PdfReader(PDF_PATH).pages
        )
        for title in EXPECTED_TITLES:
            self.assertEqual(1, docx_text.count(title), title)
            self.assertEqual(1, pdf_text.count(title), title)
        for token in ("TODO", "TBD", "PLACEHOLDER"):
            self.assertNotIn(token, docx_text.upper())
            self.assertNotIn(token, pdf_text.upper())


if __name__ == "__main__":
    unittest.main()
