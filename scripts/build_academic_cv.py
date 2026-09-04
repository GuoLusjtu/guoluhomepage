"""Build Guo Lu's academic CV from the reviewed Markdown content."""

from __future__ import annotations

import argparse
import html
import re
import zipfile
import unicodedata
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
CONTENT_PATH = ROOT / "docs" / "cv-content.md"
DOCX_PATH = ROOT / "files" / "Guo-Lu-CV.docx"

BLACK = RGBColor(0, 0, 0)
BLUE_RGB = RGBColor(23, 54, 93)
FIXED_TIME = (2000, 1, 1, 0, 0, 0)
PROFILE_LINKS = {
    "https://guolusjtu.github.io/guoluhomepage/",
    "https://scholar.google.com/citations?user=R9iwlJcAAAAJ&hl=en",
    "https://github.com/GuoLusjtu?tab=repositories",
    "https://www.linkedin.com/in/guo-lu-118a6592/",
}


def section_text(text: str, heading: str) -> str:
    match = re.search(
        rf"^## {re.escape(heading)}\s*$\n(.*?)(?=^## |\Z)",
        text,
        flags=re.MULTILINE | re.DOTALL,
    )
    if not match:
        raise ValueError(f"Missing section: {heading}")
    return match.group(1).strip()


def clean_fact(line: str) -> str:
    line = re.sub(r"^\s*-\s*", "", line.strip())
    line = re.sub(r"\[\^F\d+\]", "", line)
    return line.strip()


def bullets(text: str, heading: str) -> list[str]:
    return [clean_fact(line) for line in section_text(text, heading).splitlines() if line.startswith("-")]


def table_rows(text: str, heading: str) -> list[list[str]]:
    lines = [line for line in section_text(text, heading).splitlines() if line.startswith("|")]
    return [
        [cell.strip() for cell in line.strip("|").split("|")]
        for line in lines[2:]
    ]


def add_hyperlink(paragraph, label: str, url: str):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    props.append(color)
    run.append(props)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    link.append(run)
    paragraph._p.append(link)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    for item in (begin, instr, separate, value, end):
        run.append(item)


def underline_name(paragraph, text: str):
    parts = text.split("Guo Lu")
    for index, part in enumerate(parts):
        if part:
            paragraph.add_run(part)
        if index < len(parts) - 1:
            paragraph.add_run("Guo Lu").underline = True


def add_bullet(doc: Document, text: str, compact: bool = False):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1.5 if compact else 2.5)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run(text)
    return paragraph


def add_publication(doc: Document, number: int, row: list[str]):
    title, authors, venue, year, link, _kind = row
    paragraph = doc.add_paragraph(style="CV Publication")
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(f"{number}. ").bold = True
    if link != "—":
        add_hyperlink(paragraph, title, link)
    else:
        paragraph.add_run(title)
    details = doc.add_paragraph(style="CV Publication Details")
    details.paragraph_format.keep_together = True
    underline_name(details, authors)
    details.add_run(f". {venue}, {year}.").italic = True


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.02

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = BLACK
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True
    title._element.get_or_add_pPr().remove_all("w:pBdr")

    for name, size, before, after in (
        ("Heading 1", 13, 8, 3),
        ("Heading 2", 10.5, 4, 1.5),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE_RGB
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style._element.get_or_add_pPr().remove_all("w:pBdr")

    list_style = styles["List Bullet"]
    list_style.font.name = "Arial"
    list_style.font.size = Pt(9.2)
    list_style.font.color.rgb = BLACK
    list_style.paragraph_format.left_indent = Inches(0.17)
    list_style.paragraph_format.first_line_indent = Inches(-0.14)

    for name, size, italic, after in (
        ("CV Subtitle", 11, False, 1),
        ("CV Contact", 8.5, False, 4),
        ("CV Publication", 8.8, False, 0),
        ("CV Publication Details", 8.2, False, 2.2),
    ):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.italic = italic
        style.font.color.rgb = BLACK
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0


def normalize_docx(path: Path):
    temporary = path.with_suffix(".normalized.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as target:
        for name in sorted(source.namelist()):
            info = zipfile.ZipInfo(name, FIXED_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            target.writestr(info, source.read(name))
    temporary.replace(path)


def normalize_visible_text(value: str) -> str:
    ascii_text = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode()
    return re.sub(r"[^a-z0-9]+", " ", ascii_text.lower()).strip()


def visible_docx_paragraphs(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
    paragraphs = re.findall(r"<w:p\b.*?</w:p>", document_xml, flags=re.DOTALL)
    visible = []
    for paragraph in paragraphs:
        fragments = re.findall(r"<w:t(?:\s[^>]*)?>(.*?)</w:t>", paragraph, flags=re.DOTALL)
        normalized = normalize_visible_text(
            html.unescape(re.sub(r"<[^>]+>", "", "".join(fragments)))
        )
        if len(normalized) > 20:
            visible.append(normalized)
    return visible


def verify_rendered_pdf(path: Path):
    reader = PdfReader(path)
    if not 4 <= len(reader.pages) <= 6:
        raise ValueError(f"Expected 4–6 pages, found {len(reader.pages)}")
    pdf_text = normalize_visible_text(
        " ".join(page.extract_text() or "" for page in reader.pages)
    )
    missing = [item for item in visible_docx_paragraphs(DOCX_PATH) if item not in pdf_text]
    if missing:
        raise ValueError(f"Rendered PDF is stale or incomplete: {missing[0]}")
    uris = []
    for page in reader.pages:
        for annotation_ref in page.get("/Annots", []):
            action = annotation_ref.get_object().get("/A")
            if action and action.get("/URI"):
                uris.append(str(action["/URI"]))
    for url in PROFILE_LINKS:
        if uris.count(url) != 1:
            raise ValueError(f"Expected one PDF link for {url}")


def build():
    source = CONTENT_PATH.read_text(encoding="utf-8")
    publications = table_rows(source, "Selected Publications")
    if len(publications) != 21:
        raise ValueError("The reviewed content must contain exactly 21 publications")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.22)
    configure_styles(doc)
    doc.core_properties.title = "Guo Lu Academic CV"
    doc.core_properties.author = "Guo Lu"
    doc.core_properties.subject = "Academic Curriculum Vitae"
    fixed_date = datetime(2026, 9, 4, tzinfo=timezone.utc)
    doc.core_properties.created = fixed_date
    doc.core_properties.modified = fixed_date

    name = doc.add_paragraph("Guo Lu", style="Title")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(style="CV Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Tenure-track Associate Professor and Ph.D. Supervisor").bold = True
    affiliation = doc.add_paragraph(style="CV Subtitle")
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.add_run("Department of Electronic Engineering, Shanghai Jiao Tong University")
    contact = doc.add_paragraph(style="CV Contact")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("luguo2014 AT sjtu.edu.cn  |  Room 403, Building 5, SEIEE  |  ")
    for index, (label, url) in enumerate((
        ("Homepage", "https://guolusjtu.github.io/guoluhomepage/"),
        ("Google Scholar", "https://scholar.google.com/citations?user=R9iwlJcAAAAJ&hl=en"),
        ("GitHub", "https://github.com/GuoLusjtu?tab=repositories"),
        ("LinkedIn", "https://www.linkedin.com/in/guo-lu-118a6592/"),
    )):
        if index:
            contact.add_run("  |  ")
        add_hyperlink(contact, label, url)

    doc.add_heading("Research Interests", level=1)
    for item in bullets(source, "Research Interests"):
        doc.add_paragraph(item)

    doc.add_heading("Academic Appointments", level=1)
    for item in bullets(source, "Academic Appointments"):
        add_bullet(doc, item)
    doc.add_heading("Education", level=1)
    for item in bullets(source, "Education"):
        add_bullet(doc, item)
    doc.add_heading("Honors and Awards", level=1)
    for item in bullets(source, "Honors and Awards"):
        add_bullet(doc, item, compact=True)

    doc.add_heading("Selected Research Projects", level=1)
    for item in bullets(source, "Selected Research Projects"):
        add_bullet(doc, item)

    doc.add_page_break()
    doc.add_heading("Selected Publications", level=1)
    grouped: dict[int, dict[str, list[list[str]]]] = {}
    for row in publications:
        grouped.setdefault(int(row[3]), {"journal": [], "conference-main": []})[row[5]].append(row)
    number = 1
    for year in sorted(grouped, reverse=True):
        if year == 2025:
            doc.add_page_break()
        doc.add_heading(str(year), level=2)
        for kind, label in (("journal", "Journal Articles"), ("conference-main", "Conference Papers")):
            rows = grouped[year][kind]
            if not rows:
                continue
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(2)
            heading.paragraph_format.space_after = Pt(1)
            heading.paragraph_format.keep_with_next = True
            run = heading.add_run(label)
            run.bold = True
            run.font.size = Pt(8.7)
            for row in rows:
                add_publication(doc, number, row)
                number += 1

    doc.add_page_break()
    doc.add_heading("Professional Services", level=1)
    for item in bullets(source, "Professional Services"):
        add_bullet(doc, item, compact=True)
    doc.add_heading("Teaching", level=1)
    for item in bullets(source, "Teaching"):
        add_bullet(doc, item)
    doc.add_heading("Tutorials Workshops and Challenges", level=1)
    for item in bullets(source, "Tutorials, Workshops, and Challenges"):
        add_bullet(doc, item, compact=True)

    footer = section.footer.paragraphs[0]
    footer.style = doc.styles["CV Contact"]
    add_page_number(footer)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    normalize_docx(DOCX_PATH)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--verify-pdf", type=Path)
    arguments = parser.parse_args()
    if arguments.verify_pdf:
        verify_rendered_pdf(arguments.verify_pdf)
    else:
        build()
